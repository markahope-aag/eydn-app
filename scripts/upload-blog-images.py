"""Extract images from blog docx files, upload to Supabase Storage, rebuild HTML."""
from docx import Document
import os, json, urllib.request, urllib.error, re
from datetime import datetime, timedelta

# Read env
env = {}
with open(r"D:\projects\eydn-app\.env.local") as f:
    for line in f:
        line = line.strip()
        if "=" in line and not line.startswith("#"):
            k, v = line.split("=", 1)
            env[k] = v

SUPABASE_URL = env["NEXT_PUBLIC_SUPABASE_URL"]
SERVICE_KEY = env["SUPABASE_SERVICE_ROLE_KEY"]


def upload_image(bucket, path, data, content_type):
    url = f"{SUPABASE_URL}/storage/v1/object/{bucket}/{path}"
    req = urllib.request.Request(url, data=data, method="POST", headers={
        "Authorization": f"Bearer {SERVICE_KEY}",
        "apikey": SERVICE_KEY,
        "Content-Type": content_type,
    })
    try:
        urllib.request.urlopen(req)
    except urllib.error.HTTPError as e:
        if e.code == 409:
            req2 = urllib.request.Request(url, data=data, method="PUT", headers={
                "Authorization": f"Bearer {SERVICE_KEY}",
                "apikey": SERVICE_KEY,
                "Content-Type": content_type,
            })
            urllib.request.urlopen(req2)
        else:
            print(f"  Upload error {e.code}: {e.read().decode()[:100]}")
            raise
    return f"{SUPABASE_URL}/storage/v1/object/public/{bucket}/{path}"


def create_bucket():
    url = f"{SUPABASE_URL}/storage/v1/bucket"
    data = json.dumps({"id": "blog-images", "name": "blog-images", "public": True}).encode()
    req = urllib.request.Request(url, data=data, headers={
        "Authorization": f"Bearer {SERVICE_KEY}",
        "apikey": SERVICE_KEY,
        "Content-Type": "application/json",
    })
    try:
        urllib.request.urlopen(req)
        print("Created blog-images bucket")
    except urllib.error.HTTPError:
        print("blog-images bucket exists")


def docx_to_html(filepath, slug):
    doc = Document(filepath)
    title = ""
    html_parts = []
    in_list = False

    image_rels = {}
    for rel_id, rel in doc.part.rels.items():
        if "image" in rel.reltype:
            ext = rel.target_part.content_type.split("/")[-1]
            if ext == "jpeg":
                ext = "jpg"
            image_rels[rel_id] = {
                "data": rel.target_part.blob,
                "ext": ext,
                "content_type": rel.target_part.content_type,
            }

    img_counter = 0
    ns = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
    ans = "{http://schemas.openxmlformats.org/drawingml/2006/main}"
    rns = "{http://schemas.openxmlformats.org/officeDocument/2006/relationships}"

    for para in doc.paragraphs:
        text = para.text.strip()
        style = para.style.name.lower()

        # Check for inline images
        drawings = para._element.findall(f".//{ns}drawing")
        if drawings:
            for drawing in drawings:
                blips = drawing.findall(f".//{ans}blip")
                for blip in blips:
                    embed = blip.get(f"{rns}embed")
                    if embed and embed in image_rels:
                        img = image_rels[embed]
                        filename = f"{slug}-{img_counter}.{img['ext']}"
                        public_url = upload_image("blog-images", filename, img["data"], img["content_type"])
                        if in_list:
                            html_parts.append("</ul>")
                            in_list = False
                        html_parts.append(f'<img src="{public_url}" alt="{slug}" loading="lazy" />')
                        print(f"  Uploaded: {filename}")
                        img_counter += 1

        if not text:
            continue

        rich_text = ""
        for run in para.runs:
            t = run.text
            if not t:
                continue
            if run.bold and run.italic:
                rich_text += f"<strong><em>{t}</em></strong>"
            elif run.bold:
                rich_text += f"<strong>{t}</strong>"
            elif run.italic:
                rich_text += f"<em>{t}</em>"
            else:
                rich_text += t
        if not rich_text.strip():
            rich_text = text

        if "heading 1" in style and not title:
            title = text
            if in_list:
                html_parts.append("</ul>")
                in_list = False
        elif "heading 2" in style:
            if in_list:
                html_parts.append("</ul>")
                in_list = False
            html_parts.append(f"<h2>{rich_text}</h2>")
        elif "heading 3" in style:
            if in_list:
                html_parts.append("</ul>")
                in_list = False
            html_parts.append(f"<h3>{rich_text}</h3>")
        elif "list" in style:
            if not in_list:
                html_parts.append("<ul>")
                in_list = True
            html_parts.append(f"<li>{rich_text}</li>")
        else:
            if in_list:
                html_parts.append("</ul>")
                in_list = False
            html_parts.append(f"<p>{rich_text}</p>")

    if in_list:
        html_parts.append("</ul>")

    content = "\n".join(html_parts)
    excerpt = ""
    for para in doc.paragraphs:
        text = para.text.strip()
        if text and "heading" not in para.style.name.lower():
            excerpt = text[:200]
            break

    cover = None
    if img_counter > 0:
        first = image_rels[list(image_rels.keys())[0]]
        cover = f"{SUPABASE_URL}/storage/v1/object/public/blog-images/{slug}-0.{first['ext']}"

    return title, content, excerpt, cover, img_counter


# Main
create_bucket()

posts = [
    (r"G:\My Drive\Downloads\Wedding Guest List App_ The Modern Way to Plan Your Guest List & RSVPs.docx", "wedding-guest-list-app"),
    (r"G:\My Drive\Downloads\Wedding Day of Binder_ How to Build a Perfect Day-Of Wedding Planning Binder.docx", "wedding-day-of-binder"),
    (r"G:\My Drive\Downloads\How to Set a Wedding Budget in 2026 (Step-by-Step Guide).docx", "how-to-set-wedding-budget-2026"),
    (r"G:\My Drive\Downloads\Wedding Planning Tools_ The Essential Digital Toolkit for Modern Couples.docx", "wedding-planning-tools-digital-toolkit"),
    (r"G:\My Drive\Downloads\Primary Wedding Planning Checklist 2026.docx", "wedding-planning-checklist-2026"),
]

base_date = datetime(2026, 3, 24)
sql_parts = []

for i, (filepath, slug) in enumerate(posts):
    pub_date = base_date - timedelta(days=i * 3)
    print(f"\nProcessing: {slug}")

    title, content, excerpt, cover, img_count = docx_to_html(filepath, slug)
    word_count = len(re.sub("<[^>]+>", "", content).split())
    read_time = max(1, round(word_count / 200))

    esc = lambda s: s.replace("'", "''")  # noqa
    cover_sql = f"'{cover}'" if cover else "NULL"

    sql_parts.append(
        f"UPDATE public.blog_posts SET\n"
        f"  content = '{esc(content)}',\n"
        f"  excerpt = '{esc(excerpt)}',\n"
        f"  cover_image = {cover_sql},\n"
        f"  read_time_minutes = {read_time},\n"
        f"  published_at = '{pub_date.isoformat()}'\n"
        f"WHERE slug = '{slug}';"
    )

    print(f"  {img_count} images, {word_count} words, {read_time} min, pub {pub_date.strftime('%b %d')}")

migration = "-- Fix blog: images from Supabase Storage + staggered dates\n" + "\n\n".join(sql_parts) + "\n"
with open(r"D:\projects\eydn-app\supabase\migrations\20260325000000_fix_blog_images_dates.sql", "w", encoding="utf-8") as f:
    f.write(migration)

print(f"\nMigration created with {len(sql_parts)} updates")
